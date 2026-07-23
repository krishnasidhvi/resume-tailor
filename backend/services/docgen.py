import io
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

def _parse_markdown_to_blocks(md_content: str):
    """A very basic parser to split md into lines for doc generation."""
    lines = md_content.split('\n')
    blocks = []
    for line in lines:
        line = line.strip()
        if not line:
            continue
        blocks.append(line)
    return blocks

def generate_docx(md_content: str) -> bytes:
    doc = Document()
    blocks = _parse_markdown_to_blocks(md_content)
    
    for block in blocks:
        if block.startswith('# '):
            doc.add_heading(block[2:], level=1)
        elif block.startswith('## '):
            doc.add_heading(block[3:], level=2)
        elif block.startswith('### '):
            doc.add_heading(block[4:], level=3)
        elif block.startswith('- ') or block.startswith('* '):
            p = doc.add_paragraph(block[2:], style='List Bullet')
        else:
            doc.add_paragraph(block)
            
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

def generate_pdf(md_content: str) -> bytes:
    out = io.BytesIO()
    doc = SimpleDocTemplate(out, pagesize=letter,
                            rightMargin=inch, leftMargin=inch,
                            topMargin=inch, bottomMargin=inch)
    styles = getSampleStyleSheet()
    Story = []
    
    blocks = _parse_markdown_to_blocks(md_content)
    for block in blocks:
        if block.startswith('# '):
            Story.append(Paragraph(block[2:], styles['Heading1']))
        elif block.startswith('## '):
            Story.append(Paragraph(block[3:], styles['Heading2']))
        elif block.startswith('### '):
            Story.append(Paragraph(block[4:], styles['Heading3']))
        elif block.startswith('- ') or block.startswith('* '):
            Story.append(Paragraph("• " + block[2:], styles['Normal']))
        else:
            Story.append(Paragraph(block, styles['Normal']))
        Story.append(Spacer(1, 0.1*inch))
        
    doc.build(Story)
    return out.getvalue()
